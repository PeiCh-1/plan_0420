import sys
import os
from docx import Document
from docx.shared import RGBColor

def patch_curriculum(path):
    if not os.path.exists(path):
        print(f"Error: Template not found at {path}")
        return
    
    doc = Document(path)
    
    # 檢查是否已經有學期標籤 (等冪性保護)
    full_text = "".join([p.text for p in doc.paragraphs])
    
    # 1. 注入學期標題標籤
    if "{#isFirstSemester}" not in full_text:
        for p in doc.paragraphs:
            if "第一學期" in p.text and "特殊教育課程計畫" in p.text:
                p.insert_paragraph_before("{#isFirstSemester}")
                break

    if "{#isSecondSemester}" not in full_text:
        for p in doc.paragraphs:
            if "第二學期" in p.text and "特殊教育課程計畫" in p.text:
                # 在第二學期標題前插入：先結束第一學期，再開始第二學期
                p.insert_paragraph_before("{/isFirstSemester}")
                p.insert_paragraph_before("{#isSecondSemester}")
                break
    
    # 文件末端閉合第二學期
    if "{/isSecondSemester}" not in full_text:
        doc.add_paragraph("{/isSecondSemester}")

    # 2. 定標填表教師 (如果範本中有此文字)
    for p in doc.paragraphs:
        if ('填表教師：' in p.text or '設計者/教學者' in p.text) and '{Teacher}' not in p.text:
             p.add_run('{Teacher}')

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if '設計者/教學者' in c.text and '{Teacher}' not in c.text:
                    c.paragraphs[0].add_run('：{Teacher}')

    # 3. 樣式注入 (週次動態循環與日期標註)
    # 我們將表格改造成動態 Row Loop: 刪除 Row 6-25，將 Row 5 設為循環起點
    for t in doc.tables:
        if len(t.rows) >= 25:
             # 從底層刪除多餘的靜態列 (25 號到 6 號)
             for i in range(25, 5, -1):
                tr = t.rows[i]._tr
                t._tbl.remove(tr)
             
             row = t.rows[5]
             
             # Cell 0: 週次與標題標記
             c0 = row.cells[0]
             c0.text = "{#Weeks}{WeekLabel}"
             
             # Cell 1: 學習表現 (IndRuns)
             c1 = row.cells[1]
             c1.text = ""
             p1 = c1.paragraphs[0]
             p1.add_run("{#IndRuns}{#isAdd}")
             r_add = p1.add_run("{text}")
             r_add.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
             p1.add_run("{/isAdd}{#isDel}")
             r_del = p1.add_run("{text}")
             r_del.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
             r_del.font.strike = True
             p1.add_run("{/isDel}{^isAdd}{^isDel}{text}{/isDel}{/isAdd}{/IndRuns}")
             
             # Cell 2: 教學重點
             row.cells[2].text = "{LessonFocus}"
             
             # Cell 3: 評量方式
             row.cells[3].text = "{Assessment}"
             
             # Cell 5: 融入議題
             row.cells[5].text = "{Issues}"
             
             # Cell 7: 備註與循環終止
             row.cells[7].text = "{Notes}{/Weeks}"
            
    # 4. 行政表頭標籤注入 (Row 0-3)
    for t in doc.tables:
        if len(t.rows) >= 25:
             # Row 0: 領域/科目 (Cell 1), 課程名稱 (Cell 6)
             r0 = t.rows[0]
             if "{DomainModeString}" not in r0.cells[1].text:
                 r0.cells[1].text = "{DomainModeString}"
             if "{CourseName}" not in r0.cells[6].text:
                 r0.cells[6].text = "{CourseName}"
             
             # Row 1: 年級 (Cell 1), 教材來源 (Cell 6)
             r1 = t.rows[1]
             if "{Grade}" not in r1.cells[1].text:
                 r1.cells[1].text = "{Grade}"
             if "{MaterialSource}" not in r1.cells[6].text:
                 r1.cells[6].text = "{MaterialSource}"
             
             # Row 2: 節數 (Cell 1), 設計者 (Cell 6)
             r2 = t.rows[2]
             if "{WeeklyPeriods}" not in r2.cells[1].text:
                 r2.cells[1].text = "{WeeklyPeriods}"
             if "{Teacher}" not in r2.cells[6].text:
                 r2.cells[6].text = "{Teacher}"
             
             # Row 3: 核心素養 (Cell 1)
             r3 = t.rows[3]
             if "{CoreCompetencies}" not in r3.cells[1].text:
                 r3.cells[1].text = "{CoreCompetencies}"

    doc.save(path)
    print(f"Patched Curriculum Template (Dynamic Loop + Header): {path}")

def patch_igp(path):
    if not os.path.exists(path):
        return
    doc = Document(path)
    
    # 注入表頭標籤 (Row 0: 課程名稱, 類型, 教師)
    for t in doc.tables:
        if len(t.rows) >= 1 and len(t.columns) >= 5:
            r0 = t.rows[0]
            if "{CourseName}" not in r0.cells[1].text:
                r0.cells[1].text = "{CourseName}"
            if "{CourseType}" not in r0.cells[3].text:
                r0.cells[3].text = "{CourseType}"
            if "{Teacher}" not in r0.cells[5].text:
                r0.cells[5].text = "{Teacher}"

    for t in doc.tables:
        if len(t.rows) >= 3:
            row2 = t.rows[2]
            cell = row2.cells[1]
            if "{#IndRuns}" not in cell.text:
                cell.text = ""
                p = cell.paragraphs[0]
                p.add_run("{#IndRuns}{#isAdd}")
                r_add = p.add_run("{text}")
                r_add.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                p.add_run("{/isAdd}{#isDel}")
                r_del = p.add_run("{text}")
                r_del.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                r_del.font.strike = True
                p.add_run("{/isDel}{^isAdd}{^isDel}{text}{/isDel}{/isAdd}{/IndRuns}")
    doc.save(path)
    print(f"Patched IGP Template (Full Header): {path}")

if __name__ == "__main__":
    patch_curriculum('public/curriculum_template.docx')
    patch_igp('public/igp_template.docx')
